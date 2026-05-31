from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Ahoum_Sessions_Marketplace_Project_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_hyperlink(cell, text, url):
    cell.text = ""
    p = cell.paragraphs[0]
    add_hyperlink(p, text, url)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "111827")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()
    return table


def add_front_page_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Field", bold=True, color="FFFFFF")
    set_cell_text(hdr[1], "Details", bold=True, color="FFFFFF")
    set_cell_shading(hdr[0], "111827")
    set_cell_shading(hdr[1], "111827")

    rows = [
        ("GitHub Repository", "https://github.com/GoutamVaishnav/fullstack", True),
        ("Frontend Deployment", "https://fullstack-oz4d0623g-goutams-projects-2519f78f.vercel.app/", True),
        ("Architecture", "Microservices", False),
        ("Frontend", "React, Vite, Tailwind CSS", False),
        ("Backend", "Django, Django REST Framework", False),
        ("Database", "PostgreSQL", False),
        ("Deployment", "Vercel frontend and Render backend", False),
    ]
    for label, value, is_link in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        if is_link:
            set_cell_hyperlink(cells[1], value, value)
        else:
            set_cell_text(cells[1], value)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_para(doc, text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Ahoum Sessions Marketplace")
run.font.name = "Arial"
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(17, 24, 39)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Detailed Full-Stack Project Report")
run.font.name = "Arial"
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(75, 85, 99)

add_front_page_table(doc)

add_heading(doc, "1. Project Overview")
add_para(
    doc,
    "Ahoum Sessions Marketplace is a full-stack session booking marketplace. It allows users to discover and book sessions, creators to publish and manage sessions, and admins to control users, sessions, and bookings from a central dashboard.",
)
add_para(
    doc,
    "The application follows a microservices architecture with separate backend services for authentication, session management, and booking management. The frontend is a React single-page application.",
)

add_heading(doc, "2. Objective")
add_bullets(
    doc,
    [
        "Build a session marketplace with user, creator, and admin roles.",
        "Allow users to browse, view, and book sessions.",
        "Allow creators to create, update, delete, and monitor sessions.",
        "Allow admins to manage users, sessions, and bookings.",
        "Provide a Docker-based local setup and cloud deployment support.",
    ],
)

add_heading(doc, "3. Complete Technology Stack")
add_heading(doc, "3.1 Frontend Tech Stack", 2)
add_table(
    doc,
    ["Technology", "Purpose"],
    [
        ["React", "Builds the single-page user interface using reusable components."],
        ["Vite", "Provides fast frontend development server and production build tooling."],
        ["JavaScript", "Main programming language used for frontend logic."],
        ["Tailwind CSS", "Utility-first CSS framework used for responsive styling and dashboard UI."],
        ["HTML", "Base document structure for the React app."],
        ["CSS", "Custom styles and Tailwind-generated styling."],
        ["Local Storage", "Stores authentication payload on the client side."],
        ["Fetch API", "Used for frontend-to-backend API communication."],
        ["Vercel", "Frontend hosting and deployment platform."],
    ],
)

add_heading(doc, "3.2 Backend Tech Stack", 2)
add_table(
    doc,
    ["Technology", "Purpose"],
    [
        ["Python", "Backend programming language."],
        ["Django", "Backend web framework used for service structure, models, and routing."],
        ["Django REST Framework", "Used to build REST APIs for auth, sessions, and bookings."],
        ["Gunicorn", "Production WSGI server used to run Django services."],
        ["PyJWT", "Used for JWT token creation and validation."],
        ["Requests", "Used for Google OAuth and service-to-service API calls."],
        ["django-cors-headers", "Handles CORS between frontend and backend services."],
        ["psycopg2-binary", "PostgreSQL database driver for Django."],
        ["Render", "Cloud platform used to deploy backend Docker services."],
    ],
)

add_heading(doc, "3.3 Database, Infrastructure, and DevOps Stack", 2)
add_table(
    doc,
    ["Technology", "Purpose"],
    [
        ["PostgreSQL", "Relational database used by backend services."],
        ["Redis", "Configured cache/message infrastructure in Docker setup."],
        ["Docker", "Containerizes frontend and backend services."],
        ["Docker Compose", "Runs all local services together with one command."],
        ["Nginx", "Acts as a local reverse proxy and serves the production frontend container."],
        ["GitHub", "Source code repository and deployment source."],
        ["Render PostgreSQL", "Managed production database."],
        ["Vercel", "Frontend deployment platform."],
    ],
)

add_heading(doc, "4. System Architecture")
add_para(doc, "Local architecture uses Docker Compose and Nginx as the single public entry point.")
add_code(
    doc,
    """
Browser
  |
  v
Nginx Reverse Proxy :8080
  |
  |-- /              -> React frontend
  |-- /api/auth/     -> Django auth-service
  |-- /api/sessions/ -> Django sessions-service
  |-- /api/bookings/ -> Django bookings-service
  |
  v
PostgreSQL + Redis
""",
)
add_para(doc, "Deployment architecture uses Vercel for frontend hosting and Render for backend services.")
add_code(
    doc,
    """
Vercel Frontend
  |
  |-- Auth API     -> Render auth-service
  |-- Sessions API -> Render sessions-service
  |-- Bookings API -> Render bookings-service
  |
  v
Render PostgreSQL
""",
)

add_heading(doc, "5. Backend Services")
add_heading(doc, "5.1 Auth Service", 2)
add_para(doc, "Folder: auth-service")
add_bullets(
    doc,
    [
        "Handles Google OAuth login and development/demo login.",
        "Generates JWT tokens after successful authentication.",
        "Stores user profile and role information.",
        "Supports user, creator, and admin roles.",
        "Provides admin APIs for user listing, role updates, blocking, unblocking, and deletion.",
    ],
)

add_heading(doc, "5.2 Sessions Service", 2)
add_para(doc, "Folder: sessions-service")
add_bullets(
    doc,
    [
        "Provides public catalog of active sessions.",
        "Supports session search by title and category.",
        "Allows creators and admins to create sessions.",
        "Allows creators to manage their own sessions.",
        "Allows admins to view and manage all sessions.",
    ],
)

add_heading(doc, "5.3 Bookings Service", 2)
add_para(doc, "Folder: bookings-service")
add_bullets(
    doc,
    [
        "Allows users to book sessions.",
        "Stores booking status as active or cancelled.",
        "Provides user booking history.",
        "Provides creator booking overview.",
        "Provides admin-level booking management.",
        "Prevents duplicate booking for the same user and session.",
    ],
)

add_heading(doc, "6. Frontend Features")
add_bullets(
    doc,
    [
        "Public catalog page for browsing sessions.",
        "Session detail page with schedule, host, price, duration, capacity, and meeting link.",
        "User dashboard with bookings and profile management.",
        "Creator dashboard with session creation, editing, deletion, and booking overview.",
        "Admin dashboard with user, session, and booking management.",
        "Role-based navigation and dashboard rendering.",
    ],
)

add_heading(doc, "7. API Endpoints")
add_heading(doc, "7.1 Auth Service API Endpoints", 2)
add_table(
    doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/api/auth/google/url/", "Returns Google OAuth login URL."],
        ["POST", "/api/auth/google/callback/", "Completes Google OAuth callback and returns JWT token."],
        ["POST", "/api/auth/dev-login/", "Creates demo login session for testing."],
        ["GET", "/api/auth/me/", "Returns current authenticated user profile."],
        ["PATCH", "/api/auth/me/", "Updates current user profile and role where allowed."],
        ["GET", "/api/auth/admin/users/", "Admin endpoint to list all users."],
        ["PATCH", "/api/auth/admin/users/", "Admin endpoint to update user role, name, or block status."],
        ["DELETE", "/api/auth/admin/users/", "Admin endpoint to delete a user."],
    ],
)

add_heading(doc, "7.2 Sessions Service API Endpoints", 2)
add_table(
    doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/api/sessions/", "Returns public list of active sessions."],
        ["POST", "/api/sessions/", "Creator/admin endpoint to create a session."],
        ["GET", "/api/sessions/mine/", "Returns sessions owned by the current creator, or all sessions for admin."],
        ["GET", "/api/sessions/admin/", "Admin endpoint to list all sessions."],
        ["GET", "/api/sessions/:id/", "Returns details of a specific session."],
        ["PATCH", "/api/sessions/:id/", "Updates a session if owner creator or admin."],
        ["DELETE", "/api/sessions/:id/", "Deletes a session if owner creator or admin."],
    ],
)

add_heading(doc, "7.3 Bookings Service API Endpoints", 2)
add_table(
    doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/api/bookings/", "Returns bookings for the current user."],
        ["POST", "/api/bookings/", "Creates a booking for a session."],
        ["GET", "/api/bookings/creator/", "Returns bookings for the current creator's sessions."],
        ["GET", "/api/bookings/admin/", "Admin endpoint to list all bookings."],
        ["PATCH", "/api/bookings/:id/", "Updates booking status, such as active or cancelled."],
        ["DELETE", "/api/bookings/:id/", "Admin endpoint to delete a booking."],
    ],
)

add_heading(doc, "8. Database Design")
add_para(doc, "The project uses PostgreSQL. In local Docker development, separate databases are created for each backend service.")
add_bullets(doc, ["auth_db for auth-service", "sessions_db for sessions-service", "bookings_db for bookings-service"])
add_para(doc, "In free-tier deployment, a single Render PostgreSQL database can also be shared by all services by pointing the service-specific database name variables to the same database.")

add_heading(doc, "9. Docker and Nginx Setup")
add_heading(doc, "9.1 Docker", 2)
add_para(doc, "Each service has its own Dockerfile. Docker is used to containerize the React frontend and all Django backend services.")
add_bullets(
    doc,
    [
        "Frontend Dockerfile builds the Vite app and serves the build output using Nginx.",
        "Each backend Dockerfile installs Python dependencies and exposes port 8000.",
        "Docker Compose starts PostgreSQL, Redis, auth-service, sessions-service, bookings-service, frontend, and nginx together.",
    ],
)
add_code(doc, "docker compose up --build")

add_heading(doc, "9.2 Nginx", 2)
add_para(doc, "Nginx acts as the local reverse proxy. It provides one public entry point and forwards requests to the correct service.")
add_table(
    doc,
    ["Path", "Target"],
    [
        ["/", "React frontend"],
        ["/api/auth/", "auth-service"],
        ["/api/sessions/", "sessions-service"],
        ["/api/bookings/", "bookings-service"],
    ],
)
add_para(doc, "The frontend container also uses Nginx to serve the production build from the dist folder.")

add_heading(doc, "10. Local Run Instructions")
add_bullets(
    doc,
    [
        "Open the project root folder in VS Code.",
        "Ensure Docker Desktop is running.",
        "Run docker compose up --build.",
        "Open http://localhost:8080 in the browser.",
        "Use Ctrl+C and docker compose down to stop the project.",
    ],
)

add_heading(doc, "11. Deployment Summary")
add_table(
    doc,
    ["Component", "Platform"],
    [
        ["Frontend", "Vercel"],
        ["auth-service", "Render Docker Web Service"],
        ["sessions-service", "Render Docker Web Service"],
        ["bookings-service", "Render Docker Web Service"],
        ["Database", "Render PostgreSQL"],
    ],
)

add_heading(doc, "12. Key Features")
add_bullets(
    doc,
    [
        "Role-based access for users, creators, and admins.",
        "Session catalog and session detail page.",
        "Creator session management.",
        "User booking management.",
        "Admin dashboard for platform control.",
        "JWT-based authentication and authorization.",
        "Dockerized local development.",
        "Cloud deployment support through Render and Vercel.",
    ],
)

add_heading(doc, "13. Limitations")
add_bullets(
    doc,
    [
        "Payment integration is not implemented.",
        "File upload is not implemented; image URLs are used.",
        "Email notification system is not implemented.",
        "Google OAuth requires correct production credentials and redirect URI configuration.",
        "Automated tests are not included yet.",
    ],
)

add_heading(doc, "14. Future Scope")
add_bullets(
    doc,
    [
        "Add Stripe or Razorpay payment integration.",
        "Add email confirmations and reminders.",
        "Add real image upload using cloud storage.",
        "Add reviews and ratings.",
        "Add automated backend and frontend tests.",
        "Improve analytics for admins and creators.",
        "Add production monitoring and logging.",
    ],
)

add_heading(doc, "15. Conclusion")
add_para(
    doc,
    "Ahoum Sessions Marketplace is a full-stack microservices-based application that demonstrates session discovery, publishing, booking, and administration. The system uses React for the frontend, Django services for backend APIs, PostgreSQL for persistence, Docker for local development, Nginx for routing, Render for backend deployment, and Vercel for frontend deployment.",
)
add_para(
    doc,
    "The project provides a strong foundation for a production-ready marketplace and can be expanded with payments, notifications, uploads, analytics, and automated testing.",
)

doc.save(OUT)
print(OUT)
